"""
NexusAI — Document Processor Service
Extracts and cleans text from PDF and DOCX files.
"""

import logging
import re
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger("nexusai.processor")


class DocumentProcessor:
    """Handles text extraction from supported document formats."""

    def extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text from a document based on its file extension.

        Args:
            file_path: Path to the document file.
            file_extension: File extension (e.g., '.pdf', '.docx').

        Returns:
            Extracted and cleaned text content.
        """
        ext = file_extension.lower().replace(".", "")
        if ext == "pdf":
            return self._extract_from_pdf(file_path)
        elif ext == "docx":
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file page by page."""
        logger.info(f"📖 Extracting text from PDF: {file_path}")
        text_parts = []

        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    cleaned = self._clean_text(page_text)
                    if cleaned:
                        text_parts.append(f"[Page {page_num + 1}]\n{cleaned}")

            total_pages = len(reader.pages)
            logger.info(f"✅ Extracted text from {total_pages} pages")

        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {str(e)}")
            raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

        return "\n\n".join(text_parts)

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        logger.info(f"📖 Extracting text from DOCX: {file_path}")
        text_parts = []

        try:
            doc = DocxDocument(file_path)

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    cleaned = self._clean_text(text)
                    if cleaned:
                        text_parts.append(cleaned)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            logger.info(f"✅ Extracted {len(text_parts)} text blocks from DOCX")

        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {str(e)}")
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")

        return "\n\n".join(text_parts)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text.

        - Remove excessive whitespace
        - Normalize line breaks
        - Remove non-printable characters
        - Strip leading/trailing whitespace
        """
        text = re.sub(r"[^\S\n]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[^\x20-\x7E\n\r\t]", "", text)
        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
        text = text.strip()
        return text
